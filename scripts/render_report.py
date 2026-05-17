#!/usr/bin/env python3
"""
render_report.py
================
Verification 결과 (per-reference verdict + field diffs)를 사람이 읽기 좋은
Markdown 리포트와 .docx 리포트로 렌더링한다.

입력 JSON schema (verifications.json):
[
  {
    "idx": 1,
    "raw": "원본 reference 문자열",
    "verdict": "verified|partial_mismatch|hallucinated|unverifiable",
    "field_diffs": {
      "pmid":   {"user": "...", "ground_truth": "...", "match": true|false},
      "authors":{"user": "...", "ground_truth": "...", "match": true|false},
      "title":  {"user": "...", "ground_truth": "...", "match": true|false},
      "journal":{"user": "...", "ground_truth": "...", "match": true|false},
      "year":   {"user": "...", "ground_truth": "...", "match": true|false},
      "volume": {"user": "...", "ground_truth": "...", "match": true|false},
      "pages":  {"user": "...", "ground_truth": "...", "match": true|false}
    },
    "claim_support": "supported|partially_supported|not_in_abstract|contradicted|null",
    "corrected_citation_vancouver": "...",
    "notes": "..."
  },
  ...
]

CLI:
  python render_report.py verifications.json --md report.md
  python render_report.py verifications.json --docx report.docx
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

VERDICT_BADGE = {
    "verified": "[VERIFIED]",
    "partial_mismatch": "[PARTIAL MISMATCH]",
    "hallucinated": "[HALLUCINATED]",
    "unverifiable": "[UNVERIFIABLE]",
}

VERDICT_DESC = {
    "verified": "PubMed에서 일치하는 문헌이 확인되었고 모든 필드가 일치합니다.",
    "partial_mismatch": "PubMed에서 문헌은 확인되었으나 일부 필드(저자/제목/연도/저널/권/페이지)가 다릅니다.",
    "hallucinated": "해당 PMID/DOI 또는 제목+저자 조합으로 PubMed에서 문헌을 찾을 수 없습니다. Hallucination 가능성이 높습니다.",
    "unverifiable": "도구 호출이 실패했거나 식별 정보가 부족하여 검증할 수 없었습니다.",
}


# --------------------------------------------------------------------------- #
# Markdown                                                                    #
# --------------------------------------------------------------------------- #


def field_table_md(field_diffs: dict[str, Any]) -> str:
    rows = ["| Field | User | PubMed (ground truth) | Match |", "|---|---|---|---|"]
    for fname, d in field_diffs.items():
        if d is None:
            continue
        u = (d.get("user") or "—").replace("|", "\\|")
        g = (d.get("ground_truth") or "—").replace("|", "\\|")
        mk = "OK" if d.get("match") else "DIFF"
        rows.append(f"| {fname} | {u} | {g} | {mk} |")
    return "\n".join(rows)


def render_md(verifications: list[dict[str, Any]]) -> str:
    counts: dict[str, int] = {k: 0 for k in VERDICT_BADGE}
    for v in verifications:
        counts[v.get("verdict", "unverifiable")] = counts.get(v.get("verdict", "unverifiable"), 0) + 1

    lines = [
        "# Reference Verification Report",
        "",
        "본 리포트는 Chain-of-Verification (CoVe; Dhuliawala et al., ACL Findings 2024) 기법의 ",
        "Factor+Revise 변형을 적용해 작성되었습니다. 각 reference의 PMID/DOI/저자/제목/저널/연도/권/페이지 ",
        "atomic field들이 PubMed 도구 호출의 직접 반환값과 비교 검증되었습니다.",
        "",
        "## Summary",
        "",
        f"- 총 검증 항목: **{len(verifications)}**",
        f"- Verified: **{counts.get('verified', 0)}**",
        f"- Partial mismatch: **{counts.get('partial_mismatch', 0)}**",
        f"- Hallucinated: **{counts.get('hallucinated', 0)}**",
        f"- Unverifiable: **{counts.get('unverifiable', 0)}**",
        "",
        "## Verdict 정의",
        "",
    ]
    for v, desc in VERDICT_DESC.items():
        lines.append(f"- **{VERDICT_BADGE[v]}** — {desc}")
    lines += ["", "---", "", "## Per-reference details", ""]

    for v in verifications:
        idx = v.get("idx", "?")
        verdict = v.get("verdict", "unverifiable")
        badge = VERDICT_BADGE.get(verdict, "[?]")
        lines += [
            f"### {idx}. {badge}",
            "",
            f"**Original (as written):**",
            "",
            f"> {v.get('raw', '').strip()}",
            "",
            "**Field-level comparison:**",
            "",
            field_table_md(v.get("field_diffs", {})),
            "",
        ]
        cs = v.get("claim_support")
        if cs:
            lines += [f"**In-text claim support (abstract-based):** `{cs}`", ""]
        if v.get("corrected_citation_vancouver"):
            lines += [
                "**Suggested corrected citation (Vancouver):**",
                "",
                f"> {v['corrected_citation_vancouver']}",
                "",
            ]
        if v.get("notes"):
            lines += [f"**Notes:** {v['notes']}", ""]
        lines += ["---", ""]

    lines += [
        "## Methodology note",
        "",
        "본 검증의 한계 (CoVe 논문 Limitations 그대로):",
        "1. CoVe는 hallucination을 완전히 제거하지 않습니다 — 감소시킬 뿐입니다.",
        "2. PubMed에 색인되지 않은 문헌(일부 conference proceedings, preprint, 단행본 등)은 검증되지 않습니다.",
        "3. Claim support 평가는 abstract만을 근거로 하며, full-text가 필요한 세부 주장은 'not_in_abstract'로 표시됩니다.",
        "",
        "이 리포트는 **사람의 최종 검토를 대체하지 않습니다**. 모든 'partial_mismatch' / 'hallucinated' 항목은 사용자가 직접 원본 논문을 확인하시기 바랍니다.",
    ]
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# DOCX                                                                        #
# --------------------------------------------------------------------------- #


def render_docx(verifications: list[dict[str, Any]], out_path: Path) -> None:
    try:
        import docx  # python-docx
        from docx.shared import Pt
    except ImportError as e:
        raise SystemExit(
            "python-docx가 필요합니다. "
            "`pip install --break-system-packages python-docx` 후 재시도하세요."
        ) from e

    doc = docx.Document()
    doc.add_heading("Reference Verification Report", level=0)

    intro = doc.add_paragraph()
    intro.add_run(
        "본 리포트는 Chain-of-Verification (Dhuliawala et al., ACL Findings 2024) "
        "Factor+Revise 변형을 사용하여 PubMed 도구 호출의 직접 반환값과 사용자의 인용을 "
        "비교한 결과입니다."
    )

    counts: dict[str, int] = {k: 0 for k in VERDICT_BADGE}
    for v in verifications:
        counts[v.get("verdict", "unverifiable")] = counts.get(v.get("verdict", "unverifiable"), 0) + 1

    doc.add_heading("Summary", level=1)
    summary = doc.add_paragraph()
    summary.add_run(f"총 {len(verifications)}개 reference 검증됨\n").bold = True
    for v_key, badge in VERDICT_BADGE.items():
        summary.add_run(f"  - {badge}: {counts.get(v_key, 0)}\n")

    for v in verifications:
        idx = v.get("idx", "?")
        verdict = v.get("verdict", "unverifiable")
        doc.add_heading(f"{idx}. {VERDICT_BADGE.get(verdict, '[?]')}", level=2)

        p = doc.add_paragraph()
        p.add_run("Original: ").bold = True
        p.add_run(v.get("raw", "").strip())

        # Field table
        diffs = v.get("field_diffs") or {}
        if diffs:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Field"
            hdr[1].text = "User"
            hdr[2].text = "PubMed"
            hdr[3].text = "Match"
            for fname, d in diffs.items():
                if d is None:
                    continue
                row = table.add_row().cells
                row[0].text = fname
                row[1].text = str(d.get("user") or "—")
                row[2].text = str(d.get("ground_truth") or "—")
                row[3].text = "OK" if d.get("match") else "DIFF"

        cs = v.get("claim_support")
        if cs:
            p2 = doc.add_paragraph()
            p2.add_run("In-text claim support (abstract): ").bold = True
            p2.add_run(cs)

        if v.get("corrected_citation_vancouver"):
            p3 = doc.add_paragraph()
            p3.add_run("Suggested corrected citation (Vancouver): ").bold = True
            p3.add_run(v["corrected_citation_vancouver"])

        if v.get("notes"):
            p4 = doc.add_paragraph()
            p4.add_run("Notes: ").bold = True
            p4.add_run(v["notes"])

    doc.add_heading("Limitations", level=1)
    lim = doc.add_paragraph()
    lim.add_run(
        "이 리포트는 사람의 최종 검토를 대체하지 않습니다. CoVe는 hallucination을 감소시키지만 "
        "완전히 제거하지 못합니다. PubMed에 색인되지 않은 문헌은 검증되지 않으며, "
        "claim_support는 abstract만을 근거로 하므로 'not_in_abstract'인 경우 full-text 확인이 필요합니다."
    )

    doc.save(str(out_path))


# --------------------------------------------------------------------------- #
# CLI                                                                         #
# --------------------------------------------------------------------------- #


def main() -> int:
    parser = argparse.ArgumentParser(description="Render CoVe verification report.")
    parser.add_argument("verifications_json", help="Path to verifications JSON.")
    parser.add_argument("--md", help="Output Markdown path.")
    parser.add_argument("--docx", help="Output .docx path.")
    args = parser.parse_args()

    if not args.md and not args.docx:
        parser.error("--md 또는 --docx 중 하나 이상 지정해야 합니다.")

    data = json.loads(Path(args.verifications_json).read_text(encoding="utf-8"))

    if args.md:
        Path(args.md).write_text(render_md(data), encoding="utf-8")
        print(f"Wrote Markdown report to {args.md}", file=sys.stderr)

    if args.docx:
        render_docx(data, Path(args.docx))
        print(f"Wrote DOCX report to {args.docx}", file=sys.stderr)

    return 0


if __name__ == "__main__":
    sys.exit(main())
